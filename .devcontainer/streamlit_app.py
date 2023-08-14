from collections import namedtuple
import altair as alt
import math
import pandas as pd
import seaborn as sns
import streamlit as st
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO
from streamlit.components.v1 import html
from docx import Document
from docx.shared import Inches
import base64
from streamlit.components.v1 import html

def generate_report(df):
    # Créer un document Word
    doc = Document()
    doc.add_heading("Analysis report", level=1)

    # Ajouter le nombre d'enregistrements au rapport
    num_records = len(df)
    doc.add_paragraph(f"Number of records : {num_records}")

    # Ajouter le tableau issu de df.describe() au rapport
    doc.add_heading("General infos:", level=2)
    df_describe_table = df.describe().reset_index()
    df_describe_table.columns = [""] + list(df_describe_table.columns[1:])
    table = doc.add_table(df_describe_table.shape[0]+1, df_describe_table.shape[1])
    for i in range(df_describe_table.shape[0]):
        for j in range(df_describe_table.shape[1]):
            table.cell(i+1, j).text = str(df_describe_table.values[i, j])
    
    #nul and non nulls values
    doc.add_heading("Number of Null and Non-Null Values", level=2)
    null_counts = df.isnull().sum()
    non_null_counts = df.notnull().sum()
    counts_df = pd.DataFrame({"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})
    counts_table = doc.add_table(counts_df.shape[0]+1, counts_df.shape[1])
    for i, (col, count) in enumerate(counts_df.items()):
        counts_table.cell(0, i).text = col
        for j, value in enumerate(count):
            counts_table.cell(j+1, i).text = str(value)
            
    # Filtrer les colonnes numériques pour la heatmap
    numeric_columns = df.select_dtypes(include=[float, int]).columns
    numeric_df = df[numeric_columns]
    
    # Ajouter un histogramme pour chaque colonne au rapport
    doc.add_heading("Histograms", level=2)
    for col in numeric_df.columns:
        plt.hist(df[col], bins=20)
        plt.title(col)
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format="png")
        plt.close()
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()

    # Ajouter la heatmap avec seaborn au rapport
    heatmap_df = df[numeric_columns]
    doc.add_heading("Heatmap", level=2)
    plt.figure(figsize=(8, 6))
    sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm')
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    plt.close()
    doc.add_picture(img_buffer, width=Inches(6))
    img_buffer.close()

    # Sauvegarder le rapport en Word
    #doc.save("rapport_analyse.docx")
    # Enregistrer le rapport dans un flux binaire
    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)

    # Générer le lien de téléchargement pour le rapport
    href = f"<a href='data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(report_buffer.read()).decode()}' download='data_analysis_report.docx'>Download report</a>"

    # Afficher le lien de téléchargement dans l'application Streamlit
    st.markdown(href, unsafe_allow_html=True)
    
def main():
    config()
    # Ajout d'un sélecteur de fichiers CSV
    uploaded_file = st.file_uploader("Please choose a CSV file", type=["csv"])

    if uploaded_file is not None:
        # Chargement du fichier CSV dans un DataFrame pandas
        df = pd.read_csv(uploaded_file)
        
        # Affichage du nombre d'enregistrements
        num_records = len(df)
        st.write("Records number :", num_records)
        st.write()
        
        # Affichage du tableau issu de df.describe()
        st.write("General infos")
        st.write(df.describe())
        st.write()

         # Affichage des 10 premiers lignes
        st.write("First lines")
        st.write(df.head())
        st.write()   
        
        # Obtenir le nombre de valeurs nulles et non nulles par colonne
        null_counts = df.isnull().sum()
        non_null_counts = df.notnull().sum()

        # Créer un DataFrame à partir des séries null_counts et non_null_counts
        counts_df = pd.DataFrame({"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})

        # Afficher le tableau des nombres de valeurs nulles et non nulles
        st.write("Number of null and non-null values by columns")
        st.table(counts_df)
        
        # Filtrer les colonnes numériques pour la heatmap
        numeric_columns = df.select_dtypes(include=[float, int]).columns
        numeric_df = df[numeric_columns]
        # Affichage d'un histogramme pour chaque colonne
        for col in numeric_df.columns:
            fig = px.histogram(df, x=col, nbins=20, title=col)
            st.plotly_chart(fig)

        
         # Afficher la heatmap interactive avec Plotly
        st.write("Heatmap")
        if not numeric_columns.empty:
            fig_heatmap = go.Figure(data=go.Heatmap(z=numeric_df.corr(), x=numeric_df.columns, y=numeric_df.columns, colorscale="RdBu"))
            st.plotly_chart(fig_heatmap)
        else:
            st.write("Aucune colonne numérique disponible pour la heatmap.")
      # Générer et proposer le rapport d'analyse en Word à l'utilisateur
        if st.button("Generate analysis report in Word format"):
            generate_report(df)
            st.success("analysis report has been generated, click on link to download.")
    # Ajouter le footer
        footer_html = """
            <div style="text-align: center; margin-top: 50px;">
                <p>Developped by Ollie </p>
                <p>Date: August 2023</p>
                <p><a href="https://github.com/olivvius/streamlit-example/tree/master">Github link of the app</a></p>
                <p>
                    <a style="color: #fff;" href="https://github.com/olivvius/streamlit-example/tree/master" target="_blank">
                        <img src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://pandas.pydata.org/" target="_blank">
                        <img src="https://upload.wikimedia.org/wikipedia/commons/thumb/e/ed/Pandas_logo.svg/2048px-Pandas_logo.svg.png" alt="Pandas" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://matplotlib.org/" target="_blank">
                        <img src="https://matplotlib.org/stable/_static/logo2_compressed.svg" alt="Matplotlib" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://seaborn.pydata.org/" target="_blank">
                        <img src="https://seaborn.pydata.org/_static/logo-wide-lightbg.svg" alt="Seaborn" height="30" width="120">
                    </a>
                </p>
            </div>
        """
        st.markdown(footer_html, unsafe_allow_html=True)
            
def config():
    st.set_option('deprecation.showPyplotGlobalUse', False)

     # Changer le thème Streamlit
    st.set_page_config(layout="wide", page_title="Data Analysis", page_icon=":bar_chart:")

    # Définir le fond gris foncé
    st.markdown(
        """
        <style>
            body {
                background-color: #333;
                color: #fff;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Ajouter une police de caractère sympathique
    st.markdown(
        """
        <style>
            body {
                font-family: 'Helvetica Neue', sans-serif;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Titre de l'application
    st.title("Exploratory Data analysis from a CSV file")

    

if __name__ == "__main__":
    main()
